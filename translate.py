"""
DubStudio — Translation + Localization + Cue Sheet
Stage 1e: Google Translate + LLM translation
Stage 2:  Name/place localization via CSV
Stage 1f: Excel cue sheet generation
"""

import csv
import json
import os
import re
import time
from pathlib import Path
from typing import Callable

from argus import make_client
from deep_translator import GoogleTranslator
from dotenv import load_dotenv
import openpyxl
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

load_dotenv()

ARGUS_MODEL = os.getenv("ARGUS_MODEL", "claude-sonnet-4-6")

LANG_NAMES = {
    "de": "German", "fr": "French", "it": "Italian",
    "es": "Spanish", "pt": "Portuguese",
}

LLM_BATCH_SIZE = 15
FPS = 25


# ── Main translation entry point ───────────────────────────────────────────────

def translate_all(
    segments: list,
    characters: dict,
    target_lang: str,
    existing: dict,
    log_fn: Callable,
    project_name: str = "",
) -> dict:
    translations = {k: dict(v) for k, v in existing.items()}
    google_key = f"google_{target_lang}"
    llm_key = f"llm_{target_lang}"
    lang_name = LANG_NAMES.get(target_lang, target_lang)

    # Find what still needs translation
    needs_google = [i for i in range(len(segments))
                    if google_key not in translations.get(str(i), {})]
    needs_llm = [i for i in range(len(segments))
                 if llm_key not in translations.get(str(i), {})]

    # ── Google Translate ───────────────────────────────────────────────────────
    if needs_google:
        log_fn(f"    Google Translate: {len(needs_google)} segments...")
        gt = GoogleTranslator(source="en", target=target_lang)
        ok = 0
        for i in needs_google:
            text = segments[i].get("text", "").strip()
            if not text:
                continue
            try:
                translated = gt.translate(text)
                # Fix double-encoding: deep_translator sometimes returns latin-1 bytes decoded as latin-1
                try:
                    translated = translated.encode("latin-1").decode("utf-8")
                except (UnicodeEncodeError, UnicodeDecodeError):
                    pass
                if str(i) not in translations:
                    translations[str(i)] = {}
                translations[str(i)][google_key] = translated
                ok += 1
            except Exception as exc:
                log_fn(f"    Google failed for seg {i}: {exc}")
            time.sleep(0.05)
        log_fn(f"    Google Translate done: {ok}/{len(needs_google)}")
    else:
        log_fn(f"    Google translations already complete.")

    # ── LLM Translation ────────────────────────────────────────────────────────
    if needs_llm:
        log_fn(f"    LLM ({lang_name}): {len(needs_llm)} segments in batches of {LLM_BATCH_SIZE}...")
        translations = _translate_llm(segments, characters, target_lang, translations, needs_llm, log_fn, project_name)
    else:
        log_fn(f"    LLM translations already complete.")

    return translations


LANG_TO_PROMPT = {
    "fr": "en_fr", "de": "en_de", "es": "en_es",
    "pt": "en_pt", "it": "en_it_step1", "hi": "en_hi",
}

LANG_TO_STEP2_PROMPT = {
    "it": "en_it_step2",
}

def _load_translation_prompt(target_lang: str) -> str | None:
    key = LANG_TO_PROMPT.get(target_lang)
    if not key:
        return None
    prompt_file = Path(__file__).parent / "prompts" / f"{key}.txt"
    if not prompt_file.exists():
        return None
    return prompt_file.read_text(encoding="utf-8-sig")

def _load_step2_prompt(target_lang: str) -> str | None:
    key = LANG_TO_STEP2_PROMPT.get(target_lang)
    if not key:
        return None
    prompt_file = Path(__file__).parent / "prompts" / f"{key}.txt"
    if not prompt_file.exists():
        return None
    return prompt_file.read_text(encoding="utf-8-sig")


def _translate_llm(
    segments: list,
    characters: dict,
    target_lang: str,
    translations: dict,
    indices: list,
    log_fn: Callable,
    project_name: str = "",
) -> dict:
    import cancel_flag
    client = make_client()
    llm_key = f"llm_{target_lang}"
    lang_name = LANG_NAMES.get(target_lang, target_lang)
    n_batches = (len(indices) + LLM_BATCH_SIZE - 1) // LLM_BATCH_SIZE

    system_prompt = _load_translation_prompt(target_lang)
    step2_prompt = _load_step2_prompt(target_lang)

    if step2_prompt:
        log_fn(f"    2-step pipeline active for {lang_name} (Step 1: linguistic, Step 2: Pocket FM creative).")
        step1_key = f"llm_{target_lang}_step1"
    elif system_prompt:
        log_fn(f"    Using detailed {lang_name} adaptation prompt.")
        step1_key = llm_key
    else:
        log_fn(f"    No detailed prompt for '{target_lang}' — using generic.")
        step1_key = llm_key

    # ── Step 1: Technical / linguistic translation ─────────────────────────────
    for b in range(n_batches):
        if project_name and cancel_flag.is_cancelled(project_name):
            log_fn(f"    LLM translation cancelled after {b}/{n_batches} batches.")
            break
        batch = indices[b * LLM_BATCH_SIZE:(b + 1) * LLM_BATCH_SIZE]
        log_fn(f"    {'Step 1 ' if step2_prompt else ''}LLM batch {b + 1}/{n_batches} ({len(batch)} segs)...")

        lines = []
        for idx in batch:
            seg = segments[idx]
            char = _get_char(characters, seg, idx)
            prev = segments[idx - 1]["text"][:60] if idx > 0 else ""
            nxt = segments[idx + 1]["text"][:60] if idx < len(segments) - 1 else ""
            lines.append(f"SEG_{idx}|{char}|{seg.get('text','')}|PREV:{prev}|NEXT:{nxt}")

        user_content = (
            f"Translate these English dialogue segments to {lang_name}.\n"
            f'Return ONLY a JSON object: {{"SEG_N": "<translation>", ...}}\n\n'
            "Segments (format: SEG_N|CHARACTER|TEXT|PREV:context|NEXT:context):\n"
            + "\n".join(lines)
        )

        messages = []
        if system_prompt:
            messages.append({"role": "system", "content": system_prompt})
        messages.append({"role": "user", "content": user_content})

        create_kwargs = {"model": ARGUS_MODEL, "max_tokens": 4096, "messages": messages}
        if step2_prompt:
            create_kwargs["temperature"] = 0.2

        try:
            resp = client.messages.create(**create_kwargs)
            if hasattr(resp, "usage"):
                log_fn(f"__TOKENS__:{resp.usage.input_tokens}:{resp.usage.output_tokens}")
            raw = resp.content[0].text.strip()
            m = re.search(r"\{.*\}", raw, re.DOTALL)
            if m:
                result = json.loads(m.group())
                for key, val in result.items():
                    try:
                        idx = int(key.replace("SEG_", ""))
                        if str(idx) not in translations:
                            translations[str(idx)] = {}
                        translations[str(idx)][step1_key] = str(val)
                    except ValueError:
                        pass
            else:
                log_fn(f"    Warning: no JSON found in LLM response for batch {b + 1}")
        except Exception as exc:
            log_fn(f"    LLM batch {b + 1} failed: {exc}")

        time.sleep(0.4)

    # ── Step 2: Creative / Pocket FM enhancement (Italian only) ───────────────
    if step2_prompt:
        step2_indices = [i for i in indices if step1_key in translations.get(str(i), {})]
        n2 = (len(step2_indices) + LLM_BATCH_SIZE - 1) // LLM_BATCH_SIZE
        log_fn(f"    Step 2 (creative): {len(step2_indices)} segs in {n2} batches...")

        for b in range(n2):
            if project_name and cancel_flag.is_cancelled(project_name):
                log_fn(f"    Step 2 cancelled after {b}/{n2} batches.")
                break
            batch = step2_indices[b * LLM_BATCH_SIZE:(b + 1) * LLM_BATCH_SIZE]
            log_fn(f"    Step 2 batch {b + 1}/{n2} ({len(batch)} segs)...")

            lines = []
            for idx in batch:
                step1_text = translations[str(idx)].get(step1_key, "")
                lines.append(f"SEG_{idx}: {step1_text}")

            user_content = (
                f"Enhance these Italian dialogue segments for Pocket FM audio drama style.\n"
                f'Return ONLY a JSON object: {{"SEG_N": "<enhanced_italian>", ...}}\n\n'
                "Italian segments to enhance:\n"
                + "\n".join(lines)
            )

            messages = [
                {"role": "system", "content": step2_prompt},
                {"role": "user", "content": user_content},
            ]

            try:
                resp = client.messages.create(
                    model=ARGUS_MODEL,
                    max_tokens=4096,
                    messages=messages,
                    temperature=0.5,
                )
                if hasattr(resp, "usage"):
                    log_fn(f"__TOKENS__:{resp.usage.input_tokens}:{resp.usage.output_tokens}")
                raw = resp.content[0].text.strip()
                m = re.search(r"\{.*\}", raw, re.DOTALL)
                if m:
                    result = json.loads(m.group())
                    for key, val in result.items():
                        try:
                            idx = int(key.replace("SEG_", ""))
                            if str(idx) not in translations:
                                translations[str(idx)] = {}
                            translations[str(idx)][llm_key] = str(val)
                        except ValueError:
                            pass
                else:
                    log_fn(f"    Warning: no JSON found in Step 2 response for batch {b + 1}")
            except Exception as exc:
                log_fn(f"    Step 2 batch {b + 1} failed: {exc}")

            time.sleep(0.4)

    return translations


# ── Stage 2: Localization ──────────────────────────────────────────────────────

def run_stage2_localize(project_dir: Path, csv_path: Path, target_lang: str, log_fn: Callable):
    log_fn("=== Stage 2: LOCALIZE ===")

    # Load replacements
    replacements = _load_replacements(csv_path, log_fn)
    if not replacements:
        log_fn("No valid replacements found in CSV. Stage 2 skipped.")
        return

    # Auto-detect near-match variants (e.g. "Daemon" when table says "Damon")
    raw_segs = json.loads((project_dir / "segments.json").read_text(encoding="utf-8-sig"))
    segs = raw_segs.get("merged", raw_segs) if isinstance(raw_segs, dict) else raw_segs
    seg_texts = [s.get("text", "") for s in segs]
    replacements = _expand_with_variants(replacements, seg_texts, log_fn)

    trans_path = project_dir / "translations.json"
    if not trans_path.exists():
        raise RuntimeError("translations.json not found — run Stage 1 first.")

    translations = json.loads(trans_path.read_text(encoding="utf-8-sig"))
    google_key = f"google_{target_lang}"
    llm_key = f"llm_{target_lang}"
    google_local = f"google_{target_lang}_local"
    llm_local = f"llm_{target_lang}_local"

    changed = 0
    for sid, trans in translations.items():
        seg_changed = False
        if google_key in trans:
            trans[google_local] = _apply(trans[google_key], replacements)
            if trans[google_local] != trans[google_key]:
                seg_changed = True
        if llm_key in trans:
            trans[llm_local] = _apply(trans[llm_key], replacements)
            if trans[llm_local] != trans[llm_key]:
                seg_changed = True
        if seg_changed:
            changed += 1

    trans_path.write_text(json.dumps(translations, ensure_ascii=False, indent=2), encoding="utf-8")
    log_fn(f"    Localization applied — {changed} segments changed")

    # Regenerate cue sheet
    log_fn("    Regenerating cue sheet...")
    raw = json.loads((project_dir / "segments.json").read_text(encoding="utf-8-sig"))
    segs = raw.get("merged", raw) if isinstance(raw, dict) else raw
    chars = json.loads((project_dir / "characters_final.json").read_text(encoding="utf-8-sig")) if (project_dir / "characters_final.json").exists() else {}
    write_excel(project_dir, segs, translations, chars, target_lang)
    log_fn("=== Stage 2 complete ===")


def _load_replacements(csv_path: Path, log_fn: Callable) -> list[tuple[str, str]]:
    ORIG_KEYS  = {"Original", "original", "Source", "source", "Original Name", "original name"}
    LOCAL_KEYS = {"Localized", "localized", "Target", "target",
                  "Localised", "localised", "Localized Name", "localised name", "Localized name"}

    import shutil, tempfile
    rows = []
    with open(csv_path, "rb") as _f:
        _magic = _f.read(4)
    is_xlsx = _magic[:2] == b"PK" or _magic[:4] == b"\xD0\xCF\x11\xE0"
    suffix = Path(csv_path).suffix.lower()
    if is_xlsx or suffix in (".xlsx", ".xls"):
        # openpyxl validates extension — copy to a tmp .xlsx path if needed
        if suffix not in (".xlsx", ".xls"):
            _tmp = tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False)
            _tmp.close()
            shutil.copy2(csv_path, _tmp.name)
            _load_path = _tmp.name
        else:
            _load_path = str(csv_path)
        wb = openpyxl.load_workbook(_load_path, read_only=True, data_only=True)
        ws = wb.active
        headers = [str(c.value).strip() if c.value else "" for c in next(ws.iter_rows(min_row=1, max_row=1))]
        for row in ws.iter_rows(min_row=2, values_only=True):
            rows.append({headers[i]: (str(v).strip() if v is not None else "") for i, v in enumerate(row)})
        wb.close()
        if _load_path != str(csv_path):
            os.unlink(_load_path)
    else:
        with open(csv_path, encoding="utf-8-sig") as f:
            rows = list(csv.DictReader(f))

    # Detect which columns to use
    all_keys = list(rows[0].keys()) if rows else []
    orig_col  = next((k for k in all_keys if k in ORIG_KEYS),  None)
    local_col = next((k for k in all_keys if k in LOCAL_KEYS), None)

    # Fallback: skip any "Type" column, use 2nd and 3rd remaining columns
    if not orig_col or not local_col:
        non_type = [k for k in all_keys if k.lower() not in ("type", "gender", "is_new",
                    "is_deleted", "profession", "dialogue_accent", "sentences")]
        if len(non_type) >= 2:
            orig_col, local_col = non_type[0], non_type[1]
            log_fn(f"    Column auto-detected: '{orig_col}' → '{local_col}'")

    reps = []
    if orig_col and local_col:
        for row in rows:
            original  = (row.get(orig_col)  or "").strip()
            localized = (row.get(local_col) or "").strip()
            if original and localized and original != localized:
                reps.append((original, localized))

    # For multi-word names, also add a first-name-only fallback entry.
    # e.g. "Damon Blake" → "Enea Neri" also adds "Damon" → "Enea"
    # so text that uses only first names still gets replaced.
    existing_sources = {r[0].lower() for r in reps}
    fallbacks = []
    for orig, loc in reps:
        orig_parts = orig.split()
        loc_parts  = loc.split()
        if len(orig_parts) > 1 and loc_parts:
            first_orig = orig_parts[0]
            first_loc  = loc_parts[0]
            if (first_orig.lower() not in existing_sources
                    and first_orig.lower() != first_loc.lower()):
                fallbacks.append((first_orig, first_loc))
                existing_sources.add(first_orig.lower())
    reps.extend(fallbacks)

    reps.sort(key=lambda x: len(x[0]), reverse=True)
    log_fn(f"    Loaded {len(reps)} replacements ({len(fallbacks)} first-name fallbacks added)")
    return reps


def _levenshtein(a: str, b: str) -> int:
    if len(a) < len(b):
        return _levenshtein(b, a)
    if not b:
        return len(a)
    prev = list(range(len(b) + 1))
    for ca in a:
        curr = [prev[0] + 1]
        for j, cb in enumerate(b):
            curr.append(min(prev[j + 1] + 1, curr[j] + 1, prev[j] + (ca != cb)))
        prev = curr
    return prev[-1]


def _expand_with_variants(
    replacements: list[tuple[str, str]],
    seg_texts: list[str],
    log_fn: Callable,
) -> list[tuple[str, str]]:
    """
    Scan English source segments for capitalized words that differ by ≤1 edit
    from a known replacement source (e.g. 'Daemon' vs 'Damon').
    Auto-adds them so transcription misspellings are still localized.
    """
    # Collect all capitalized words that appear in the source text
    cap_words: set[str] = set()
    for text in seg_texts:
        cap_words.update(re.findall(r"\b[A-Z][a-z]+\b", text))

    existing_sources = {r[0].lower() for r in replacements}
    variants: list[tuple[str, str]] = []

    for orig, loc in replacements:
        for word in cap_words:
            if word.lower() in existing_sources:
                continue
            dist = _levenshtein(orig.lower(), word.lower())
            # Allow 1 edit for names up to 8 chars, 2 edits for longer
            max_dist = 1 if len(orig) <= 8 else 2
            if 0 < dist <= max_dist:
                variants.append((word, loc))
                existing_sources.add(word.lower())
                log_fn(f"    Variant detected: '{word}' ≈ '{orig}' → will map to '{loc}'")

    if variants:
        replacements = list(replacements) + variants
        replacements.sort(key=lambda x: len(x[0]), reverse=True)
    return replacements


def _apply(text: str, replacements: list[tuple[str, str]]) -> str:
    for original, localized in replacements:
        pattern = r"\b" + re.escape(original) + r"\b"
        text = re.sub(pattern, localized, text, flags=re.IGNORECASE)
    return text


# ── Excel cue sheet ────────────────────────────────────────────────────────────

def write_excel(
    project_dir: Path,
    segments: list,
    translations: dict,
    characters: dict,
    target_lang: str,
):
    import csv as _csv
    lang = LANG_NAMES.get(target_lang, target_lang)
    google_key = f"google_{target_lang}"
    llm_key = f"llm_{target_lang}"
    local_key = f"llm_{target_lang}_local"
    local_google_key = f"google_{target_lang}_local"

    # Load localization mapping using the same robust parser as Stage 2
    loc_map: dict[str, str] = {}
    loc_csv = project_dir / "localization.csv"
    if loc_csv.exists():
        try:
            for orig, repl in _load_replacements(loc_csv, lambda _: None):
                loc_map[orig.lower()] = repl
        except Exception:
            pass

    def localize_char(name: str) -> str:
        # Try full name, then strip _VO suffix, then title-case variant
        base = name.upper().replace("_VO", "").replace("_", " ").strip()
        return (
            loc_map.get(name.lower())
            or loc_map.get(base.lower())
            or loc_map.get(name.rstrip("_VO").lower())
            or name
        )

    # Load dub state for actual TTS durations
    dub_state: dict = {}
    dub_path = project_dir / "dub_state.json"
    if dub_path.exists():
        try:
            import json as _json
            dub_state = _json.loads(dub_path.read_text(encoding="utf-8-sig"))
        except Exception:
            pass

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Cue Sheet"

    HEADER_FILL  = PatternFill("solid", fgColor="1E1E2E")
    RED_FILL     = PatternFill("solid", fgColor="FF4444")
    AMBER_FILL   = PatternFill("solid", fgColor="FFAA00")
    GREEN_FILL   = PatternFill("solid", fgColor="1A4D2E")
    WHITE_BOLD   = Font(color="FFFFFF", bold=True)
    WRAP         = Alignment(wrap_text=True, vertical="top")
    CENTER_WRAP  = Alignment(wrap_text=True, vertical="top", horizontal="center")

    headers = [
        "ID", "Start", "End", "TC In", "TC Out",
        "Character", "Character (Localized)", "English",
        f"Google ({lang})", f"LLM ({lang})", f"Localized ({lang})",
        "Runtime (EN)", f"Runtime ({lang})",
    ]
    col_widths = [6, 9, 9, 14, 14, 18, 18, 50, 50, 50, 50, 14, 14]

    for col, (h, w) in enumerate(zip(headers, col_widths), 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = HEADER_FILL
        cell.font = WHITE_BOLD
        cell.alignment = Alignment(horizontal="center", vertical="center")
        ws.column_dimensions[get_column_letter(col)].width = w

    ws.freeze_panes = "A2"
    ws.row_dimensions[1].height = 20

    for i, seg in enumerate(segments):
        sid = str(i)
        trans = translations.get(sid, {})
        char = characters.get(str(i), "?")
        char_loc = localize_char(char)
        english = seg.get("text", "")
        google = trans.get(google_key, "")
        llm = trans.get(llm_key, "")
        localized = trans.get(local_key) or trans.get(local_google_key, "")
        tc_in = _to_tc(seg["start"])
        tc_out = _to_tc(seg["end"])
        runtime_en = round(seg["end"] - seg["start"], 2)
        dub = dub_state.get(sid, {})
        runtime_loc = round(dub["actual_duration"], 2) if dub.get("actual_duration") else None

        row = i + 2
        for col, val in enumerate(
            [i, round(seg["start"], 3), round(seg["end"], 3),
             tc_in, tc_out, char, char_loc, english, google, llm, localized,
             runtime_en, runtime_loc], 1
        ):
            cell = ws.cell(row=row, column=col, value=val)
            if col in (4, 5):  # TC In / TC Out — never wrap
                cell.alignment = Alignment(horizontal="center", vertical="top", wrap_text=False)
            elif col in (1, 2, 3, 12, 13):
                cell.alignment = CENTER_WRAP
            else:
                cell.alignment = WRAP
            if col in (12, 13) and val is not None:
                cell.number_format = "0.00"

        # Color-code runtime columns only
        if runtime_loc:
            diff = runtime_loc - runtime_en
            rt_fill = RED_FILL if diff > 0.5 else (AMBER_FILL if diff > 0.1 else GREEN_FILL)
            ws.cell(row=row, column=12).fill = rt_fill
            ws.cell(row=row, column=13).fill = rt_fill

    wb.save(project_dir / "cue_sheet_final.xlsx")


# ── Helpers ────────────────────────────────────────────────────────────────────

def _get_char(characters: dict, seg: dict, fallback_idx: int) -> str:
    return characters.get(str(fallback_idx), "?")


def _to_tc(seconds: float) -> str:
    total = int(seconds * FPS)
    frames = total % FPS
    secs = (total // FPS) % 60
    mins = (total // FPS // 60) % 60
    hours = total // FPS // 3600
    return f"{hours:02d}:{mins:02d}:{secs:02d}:{frames:02d}"


# ── Word review document ───────────────────────────────────────────────────────

def write_review_doc(
    project_dir: Path,
    segments: list,
    translations: dict,
    characters: dict,
    target_lang: str,
    project_name: str,
):
    """
    Generate a Word document (.docx) for human review after Stage 1.
    Formatted as a clean script-style document with English + translation side by side.
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import datetime

    lang = LANG_NAMES.get(target_lang, target_lang)
    llm_key = f"llm_{target_lang}"
    google_key = f"google_{target_lang}"

    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ── Title block ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("DubStudio — Stage 1 Review")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0xFF, 0x6B, 0x35)  # PocketFM orange

    meta = doc.add_paragraph()
    meta.add_run(f"Project: ").bold = True
    meta.add_run(f"{project_name}    ")
    meta.add_run(f"Language: ").bold = True
    meta.add_run(f"{lang}    ")
    meta.add_run(f"Segments: ").bold = True
    meta.add_run(f"{len(segments)}    ")
    meta.add_run(f"Generated: ").bold = True
    meta.add_run(datetime.datetime.now().strftime("%d %b %Y, %H:%M"))
    meta.runs[0].font.size = Pt(9)
    for run in meta.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x70)

    doc.add_paragraph()

    # ── Character summary ──
    from collections import Counter
    char_counts = Counter(characters.values())
    if char_counts:
        heading = doc.add_paragraph()
        h_run = heading.add_run("Characters Detected")
        h_run.bold = True
        h_run.font.size = Pt(10)
        h_run.font.color.rgb = RGBColor(0x33, 0x33, 0x55)

        char_para = doc.add_paragraph()
        char_para.paragraph_format.left_indent = Cm(0.5)
        for char, count in char_counts.most_common():
            r = char_para.add_run(f"{char}: {count} lines    ")
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x44, 0x44, 0x66)

        doc.add_paragraph()

    # ── Segment table ──
    heading2 = doc.add_paragraph()
    h2_run = heading2.add_run("Translation Review")
    h2_run.bold = True
    h2_run.font.size = Pt(10)
    h2_run.font.color.rgb = RGBColor(0x33, 0x33, 0x55)
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    headers = ["#  TC In", "Character", "English", f"Translation ({lang})"]
    col_widths_cm = [2.8, 2.6, 7.0, 7.0]
    for i, (cell, text) in enumerate(zip(hdr_cells, headers)):
        cell.text = text
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(cell, "1E1E35")
        _set_col_width(table, i, col_widths_cm[i])

    # Data rows
    for i, seg in enumerate(segments):
        sid = str(i)
        trans = translations.get(sid, {})
        char = _get_char(characters, seg, i)
        english = seg.get("text", "").strip()
        translation = trans.get(llm_key) or trans.get(google_key) or ""
        tc = _to_tc(seg["start"])
        duration = seg["end"] - seg["start"]

        row_cells = table.add_row().cells

        # Alternating row background
        bg = "F8F8FC" if i % 2 == 0 else "FFFFFF"

        # Col 0: segment # + timecode
        row_cells[0].text = f"#{i}\n{tc}\n{duration:.1f}s"
        _style_cell(row_cells[0], Pt(7.5), RGBColor(0x44, 0x44, 0x66), bg, bold_first=True)

        # Col 1: character
        row_cells[1].text = char
        _style_cell(row_cells[1], Pt(8), RGBColor(0xFF, 0x6B, 0x35), bg)

        # Col 2: English
        row_cells[2].text = english
        _style_cell(row_cells[2], Pt(8.5), RGBColor(0x1A, 0x1A, 0x2E), bg)

        # Col 3: Translation
        row_cells[3].text = translation
        color = RGBColor(0x1A, 0x1A, 0x2E) if translation else RGBColor(0xAA, 0xAA, 0xCC)
        _style_cell(row_cells[3], Pt(8.5), color, bg)
        if not translation:
            row_cells[3].paragraphs[0].runs[0].italic = True

    out_path = project_dir / f"{project_name}_stage1_review.docx"
    doc.save(out_path)
    return out_path


def write_stage3_review_doc(
    project_dir: Path,
    segments: list,
    dub_state: dict,
    target_lang: str,
    project_name: str,
) -> Path:
    """
    Generate a Word document (.docx) for human review after Stage 3 (TTS).
    Shows the final dubbed text used for each segment, TTS timing, and status.
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from collections import Counter
    import datetime

    lang = LANG_NAMES.get(target_lang, target_lang)

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("DubStudio — Stage 3 Dubbed Script")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0xFF, 0x6B, 0x35)

    tts_done = sum(1 for s in dub_state.values() if s.get("status") == "done")
    tts_err = sum(1 for s in dub_state.values() if s.get("status") == "error")
    extends = sum(1 for s in dub_state.values() if s.get("extend_by", 0) > 0)

    meta = doc.add_paragraph()
    for label, value in [
        ("Project: ", project_name + "    "),
        ("Language: ", lang + "    "),
        ("Segments: ", f"{len(segments)}    "),
        ("TTS done: ", f"{tts_done}/{len(segments)}    "),
        ("Errors: ", f"{tts_err}    "),
        ("Freeze-frames: ", f"{extends}    "),
        ("Generated: ", datetime.datetime.now().strftime("%d %b %Y, %H:%M")),
    ]:
        r = meta.add_run(label)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x44, 0x44, 0x66)
        r2 = meta.add_run(value)
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(0x55, 0x55, 0x70)

    doc.add_paragraph()

    char_counts = Counter(
        s["character"] for s in dub_state.values()
        if s.get("character") and s["character"] not in ("?", "UNKNOWN")
    )
    if char_counts:
        h = doc.add_paragraph()
        h_run = h.add_run("Characters")
        h_run.bold = True
        h_run.font.size = Pt(10)
        h_run.font.color.rgb = RGBColor(0x33, 0x33, 0x55)

        cp = doc.add_paragraph()
        cp.paragraph_format.left_indent = Cm(0.5)
        for char, count in char_counts.most_common():
            r = cp.add_run(f"{char}: {count} segments    ")
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x44, 0x44, 0x66)
        doc.add_paragraph()

    h2 = doc.add_paragraph()
    h2_run = h2.add_run(f"Dubbed Script — {lang}")
    h2_run.bold = True
    h2_run.font.size = Pt(10)
    h2_run.font.color.rgb = RGBColor(0x33, 0x33, 0x55)
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    headers = ["# TC In", "Character", "English", f"Dubbed ({lang})", "TTS"]
    col_widths_cm = [2.5, 2.5, 6.0, 6.5, 1.8]
    for i, (cell, text) in enumerate(zip(hdr_cells, headers)):
        cell.text = text
        hrun = cell.paragraphs[0].runs[0]
        hrun.bold = True
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(cell, "1E1E35")
        _set_col_width(table, i, col_widths_cm[i])

    for i, seg in enumerate(segments):
        sid = str(i)
        state = dub_state.get(sid, {})
        status = state.get("status", "pending")

        english = state.get("english") or seg.get("text", "").strip()
        dubbed = state.get("translated_text", "")
        char = state.get("character") or "?"
        tc = _to_tc(seg["start"])
        target_dur = seg["end"] - seg["start"]
        actual_dur = state.get("actual_duration")
        extend_by = state.get("extend_by", 0)
        iterations = state.get("iterations", 0)

        row_cells = table.add_row().cells
        bg = "F8F8FC" if i % 2 == 0 else "FFFFFF"

        row_cells[0].text = f"#{i}\n{tc}\n{target_dur:.1f}s"
        _style_cell(row_cells[0], Pt(7.5), RGBColor(0x44, 0x44, 0x66), bg, bold_first=True)

        row_cells[1].text = char
        _style_cell(row_cells[1], Pt(8), RGBColor(0xFF, 0x6B, 0x35), bg)

        row_cells[2].text = english
        _style_cell(row_cells[2], Pt(8.5), RGBColor(0x1A, 0x1A, 0x2E), bg)

        row_cells[3].text = dubbed if dubbed else "(no translation)"
        dubbed_color = RGBColor(0x1A, 0x1A, 0x2E) if dubbed else RGBColor(0xAA, 0xAA, 0xCC)
        _style_cell(row_cells[3], Pt(8.5), dubbed_color, bg)
        if not dubbed:
            row_cells[3].paragraphs[0].runs[0].italic = True

        if status == "done" and actual_dur is not None:
            diff = actual_dur - target_dur
            if extend_by > 0:
                tts_text = f"{actual_dur:.1f}s\n+{extend_by:.1f}s ext"
                tts_bg, tts_color = "FFF3E0", RGBColor(0xE6, 0x5C, 0x00)
            elif abs(diff) <= 0.5:
                tts_text = f"{actual_dur:.1f}s"
                tts_bg, tts_color = "E8F5E9", RGBColor(0x2E, 0x7D, 0x32)
            else:
                tts_text = f"{actual_dur:.1f}s"
                tts_bg, tts_color = "FFF8E1", RGBColor(0xB8, 0x7A, 0x00)
            if iterations > 1:
                tts_text += f"\n×{iterations} iter"
            row_cells[4].text = tts_text
            _style_cell(row_cells[4], Pt(7.5), tts_color, tts_bg)
        elif status == "error":
            row_cells[4].text = "ERR"
            _style_cell(row_cells[4], Pt(8), RGBColor(0xC6, 0x28, 0x28), "FFEBEE")
        else:
            row_cells[4].text = "—"
            _style_cell(row_cells[4], Pt(8), RGBColor(0x99, 0x99, 0xAA), bg)

    out_path = project_dir / f"{project_name}_stage3_dubbed.docx"
    doc.save(out_path)
    return out_path


def _style_cell(cell, font_size, font_color, bg_hex: str, bold_first=False):
    from docx.shared import Pt
    _set_cell_bg(cell, bg_hex)
    for i, para in enumerate(cell.paragraphs):
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        for run in para.runs:
            run.font.size = font_size
            run.font.color.rgb = font_color
            if bold_first and i == 0:
                run.bold = True


def _set_cell_bg(cell, hex_color: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_col_width(table, col_idx: int, width_cm: float):
    from docx.shared import Cm
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)
